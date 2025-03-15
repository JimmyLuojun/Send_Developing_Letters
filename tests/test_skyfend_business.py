import pytest
from docx import Document
from src.data.skyfend_business import read_skyfend_business
from unittest.mock import patch, mock_open

def test_read_skyfend_business_success(tmp_path):
    # Create a temporary .docx file for testing
    content = "Skyfend's main business is drone detection."
    doc = Document()
    doc.add_paragraph(content)
    file_path = tmp_path / "test_skyfend.docx"
    doc.save(file_path)

    # Call the function and check the result
    result = read_skyfend_business(str(file_path))
    assert result == content

def test_read_skyfend_business_file_not_found():
    # Test with a non-existent file
    result = read_skyfend_business("nonexistent_file.docx")
    assert result == ""  # Expecting an empty string for file not found

def test_read_skyfend_business_empty_file(tmp_path):
    # Test file with no text content
    doc = Document()
    file_path = tmp_path / 'empty.docx'
    doc.save(file_path)
    result = read_skyfend_business(str(file_path))
    assert result == ""

def test_read_skyfend_business_multiple_paragraphs(tmp_path):
    # Create a temp docx file with multiple paragraphs
    doc = Document()
    doc.add_paragraph("Paragraph 1")
    doc.add_paragraph("Paragraph 2")
    file_path = tmp_path / 'multi.docx'
    doc.save(file_path)
    result = read_skyfend_business(str(file_path))
    assert result == "Paragraph 1\nParagraph 2"
