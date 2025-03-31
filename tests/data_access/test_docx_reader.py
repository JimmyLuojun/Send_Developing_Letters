import pytest
from pathlib import Path
from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError
from unittest.mock import patch, MagicMock

from src.data_access.docx_reader import read_skyfend_business

# Constants
TEST_FILE_PATH = Path("dummy/path/doc.docx")
EXPECTED_TEXT = "Paragraph 1.\nParagraph 2 has text."

@pytest.fixture
def mock_docx_document():
    """Create a mock docx document with test content."""
    mock_doc = MagicMock(spec=DocxDocument)
    
    # Create paragraphs with different content types
    paragraphs = [
        MagicMock(text="Paragraph 1."),
        MagicMock(text="Paragraph 2 has text."),
        MagicMock(text=" "),  # Empty paragraph
        MagicMock(text=""),   # Empty string paragraph
        MagicMock(text=None)  # None paragraph
    ]
    
    mock_doc.paragraphs = paragraphs
    return mock_doc

@pytest.fixture
def temp_docx_file(tmp_path):
    """Create a temporary docx file for testing."""
    d = tmp_path / "data"
    d.mkdir()
    p = d / "test_doc.docx"
    p.write_text("dummy content")
    return p

def test_read_skyfend_business_success(mock_docx_document):
    """Test successful reading of a valid docx file."""
    with patch('src.data_access.docx_reader.Document', return_value=mock_docx_document) as mock_doc_init, \
         patch('src.data_access.docx_reader.Path.is_file', return_value=True):
        
        result = read_skyfend_business(TEST_FILE_PATH)
        
        # Verify the result
        assert result == EXPECTED_TEXT
        
        # Verify the Document was initialized with correct path
        mock_doc_init.assert_called_once_with(TEST_FILE_PATH)

def test_read_skyfend_business_file_not_found():
    """Test handling of non-existent file."""
    with patch('src.data_access.docx_reader.Path.is_file', return_value=False):
        result = read_skyfend_business(TEST_FILE_PATH)
        assert result is None

def test_read_skyfend_business_docx_error():
    """Test handling of docx-specific errors."""
    with patch('src.data_access.docx_reader.Document', 
              side_effect=PackageNotFoundError("Invalid or corrupted docx file")) as mock_doc_init, \
         patch('src.data_access.docx_reader.Path.is_file', return_value=True):
        
        result = read_skyfend_business(TEST_FILE_PATH)
        
        # Verify the result
        assert result is None
        
        # Verify the Document was initialized with correct path
        mock_doc_init.assert_called_once_with(TEST_FILE_PATH)

def test_read_skyfend_business_generic_error():
    """Test handling of generic errors during file reading."""
    with patch('src.data_access.docx_reader.Document', 
              side_effect=Exception("Unexpected error during file reading")) as mock_doc_init, \
         patch('src.data_access.docx_reader.Path.is_file', return_value=True):
        
        result = read_skyfend_business(TEST_FILE_PATH)
        
        # Verify the result
        assert result is None
        
        # Verify the Document was initialized with correct path
        mock_doc_init.assert_called_once_with(TEST_FILE_PATH)

def test_read_skyfend_business_empty_document():
    """Test handling of an empty document."""
    mock_empty_doc = MagicMock(spec=DocxDocument)
    mock_empty_doc.paragraphs = []
    
    with patch('src.data_access.docx_reader.Document', return_value=mock_empty_doc) as mock_doc_init, \
         patch('src.data_access.docx_reader.Path.is_file', return_value=True):
        
        result = read_skyfend_business(TEST_FILE_PATH)
        
        # Verify the result
        assert result == ""
        
        # Verify the Document was initialized with correct path
        mock_doc_init.assert_called_once_with(TEST_FILE_PATH)

def test_read_skyfend_business_with_special_characters():
    """Test handling of documents with special characters."""
    mock_special_doc = MagicMock(spec=DocxDocument)
    paragraphs = [
        MagicMock(text="Special chars: !@#$%^&*()"),
        MagicMock(text="Unicode: 你好世界"),
        MagicMock(text="Newlines: \n\t\r"),
        MagicMock(text="Quotes: \"'")
    ]
    mock_special_doc.paragraphs = paragraphs
    
    with patch('src.data_access.docx_reader.Document', return_value=mock_special_doc) as mock_doc_init, \
         patch('src.data_access.docx_reader.Path.is_file', return_value=True):
        
        result = read_skyfend_business(TEST_FILE_PATH)
        
        # Verify the result contains all paragraphs properly joined
        expected = "Special chars: !@#$%^&*()\nUnicode: 你好世界\nNewlines: \n\t\r\nQuotes: \"'"
        assert result == expected
        
        # Verify the Document was initialized with correct path
        mock_doc_init.assert_called_once_with(TEST_FILE_PATH)

def test_read_skyfend_business_with_whitespace():
    """Test handling of documents with various whitespace."""
    mock_whitespace_doc = MagicMock(spec=DocxDocument)
    paragraphs = [
        MagicMock(text="  Leading spaces"),
        MagicMock(text="Trailing spaces  "),
        MagicMock(text="  Both ends  "),
        MagicMock(text="\tTabbed\ttext"),
        MagicMock(text="\nNewlined\ntext"),
        MagicMock(text="  Multiple   spaces  ")
    ]
    mock_whitespace_doc.paragraphs = paragraphs
    
    with patch('src.data_access.docx_reader.Document', return_value=mock_whitespace_doc) as mock_doc_init, \
         patch('src.data_access.docx_reader.Path.is_file', return_value=True):
        
        result = read_skyfend_business(TEST_FILE_PATH)
        
        # Verify the result preserves whitespace within paragraphs
        expected = "  Leading spaces\nTrailing spaces  \n  Both ends  \n\tTabbed\ttext\n\nNewlined\ntext\n  Multiple   spaces  "
        assert result == expected
        
        # Verify the Document was initialized with correct path
        mock_doc_init.assert_called_once_with(TEST_FILE_PATH)

def test_read_skyfend_business_with_none_paragraphs():
    """Test handling of documents with None paragraphs."""
    mock_none_doc = MagicMock(spec=DocxDocument)
    paragraphs = [
        MagicMock(text=None),
        MagicMock(text="Valid paragraph"),
        MagicMock(text=None),
        MagicMock(text="Another valid paragraph"),
        MagicMock(text=None)
    ]
    mock_none_doc.paragraphs = paragraphs
    
    with patch('src.data_access.docx_reader.Document', return_value=mock_none_doc) as mock_doc_init, \
         patch('src.data_access.docx_reader.Path.is_file', return_value=True):
        
        result = read_skyfend_business(TEST_FILE_PATH)
        
        # Verify the result only includes valid paragraphs
        expected = "Valid paragraph\nAnother valid paragraph"
        assert result == expected
        
        # Verify the Document was initialized with correct path
        mock_doc_init.assert_called_once_with(TEST_FILE_PATH) 